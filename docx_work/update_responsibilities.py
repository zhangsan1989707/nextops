from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph


SRC = Path("/Users/leohang/Documents/My-Vault/00-inbox/成都市蓉政通运维保障方案V1.docx")
OUT = Path("/Users/leohang/Documents/My-Vault/00-inbox/成都市蓉政通运维保障方案V1-职责分工调整.docx")


def clear_paragraph(paragraph):
    p = paragraph._p
    for child in list(p):
        p.remove(child)
    return paragraph


def set_text(paragraph, text, style=None):
    clear_paragraph(paragraph)
    if style:
        paragraph.style = style
    paragraph.add_run(text)
    return paragraph


def insert_after(paragraph, text, style=None):
    new_p = deepcopy(paragraph._p)
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    return set_text(new_para, text, style)


def remove_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    paragraph._p = paragraph._element = None


def set_cell_text(cell, text, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT):
    for para in cell.paragraphs:
        clear_paragraph(para)
    para = cell.paragraphs[0]
    para.alignment = align
    run = para.add_run(text)
    run.bold = bold
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_cell_width(cell, width):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.tcW
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def set_table_fixed_width(table, widths):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[idx])


def replace_first_table(doc, rows):
    old_table = doc.tables[0]
    old_tbl = old_table._tbl
    style = old_table.style

    table = doc.add_table(rows=len(rows), cols=5)
    table.style = style

    for r_idx, row_data in enumerate(rows):
        for c_idx, text in enumerate(row_data):
            align = WD_ALIGN_PARAGRAPH.CENTER if c_idx < 4 or r_idx < 2 else WD_ALIGN_PARAGRAPH.LEFT
            set_cell_text(table.cell(r_idx, c_idx), text, bold=r_idx < 2, align=align)

    merged = table.cell(0, 0).merge(table.cell(0, 4))
    set_cell_text(merged, rows[0][0], bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_table_fixed_width(table, [1300, 1200, 1900, 1700, 3300])

    old_tbl.addprevious(table._tbl)
    old_tbl.getparent().remove(old_tbl)


doc = Document(SRC)

start = None
end = None
for i, para in enumerate(doc.paragraphs):
    if para.text.strip() == "运维组织管理":
        start = i
    elif start is not None and para.text.strip() == "运维服务内容":
        end = i
        break

if start is None or end is None:
    raise RuntimeError("未找到需要替换的“运维组织管理”章节边界")

heading2_style = doc.paragraphs[start].style
heading3_style = doc.paragraphs[start + 1].style
body_style = doc.paragraphs[start + 2].style

new_content = [
    ("主要职责分工", heading2_style),
    ("蓉政通二期分工", heading3_style),
    ("本项目按照项目管理、项目运营、项目运维三个工作面进行职责划分，形成联合管理、数城运营、中天运维的协同保障机制。", body_style),
    ("项目管理（联合）", heading3_style),
    ("人员组织：中天1人+数城1人。", body_style),
    ("主要事项：", body_style),
    ("（1）项目管理制度制定。", body_style),
    ("（2）项目对外汇报机制。", body_style),
    ("（3）管理职责分工。", body_style),
    ("项目运营（数城）", heading3_style),
    ("日常运营工作：", body_style),
    ("（1）客户现场贴身服务。", body_style),
    ("（2）项目运营台账管理。", body_style),
    ("（3）系统使用及政策执行的答疑。", body_style),
    ("（4）系统问题收集跟踪和反馈。", body_style),
    ("（5）系统需求收集跟踪和反馈。", body_style),
    ("（6）第三方平台或系统沟通协调。", body_style),
    ("（7）系统培训服务。", body_style),
    ("验收工作。", body_style),
    ("项目运维（中天）", heading3_style),
    ("日常运维工作：", body_style),
    ("（1）系统日常监控巡检。", body_style),
    ("（2）应用接入技术对接。", body_style),
    ("（3）系统故障分析和处理。", body_style),
    ("（4）系统版本升级发布。", body_style),
    ("（5）安全漏洞修复。", body_style),
    ("验收工作。", body_style),
]

set_text(doc.paragraphs[start], new_content[0][0], new_content[0][1])
current = doc.paragraphs[start]
for text, style in new_content[1:]:
    current = insert_after(current, text, style)

for para in list(doc.paragraphs)[start + len(new_content):end + len(new_content) - 1]:
    remove_paragraph(para)

table_rows = [
    ["蓉政通二期职责分工", "", "", "", ""],
    ["工作面", "责任主体", "人员组织/范围", "职责类型", "具体事项"],
    ["项目管理", "联合", "中天1人+数城1人", "人员组织", "中天1人+数城1人"],
    ["项目管理", "联合", "中天1人+数城1人", "管理事项", "项目管理制度制定"],
    ["项目管理", "联合", "中天1人+数城1人", "管理事项", "项目对外汇报机制"],
    ["项目管理", "联合", "中天1人+数城1人", "管理事项", "管理职责分工"],
    ["项目运营", "数城", "-", "日常运营工作", "客户现场贴身服务"],
    ["项目运营", "数城", "-", "日常运营工作", "项目运营台账管理"],
    ["项目运营", "数城", "-", "日常运营工作", "系统使用及政策执行的答疑"],
    ["项目运营", "数城", "-", "日常运营工作", "系统问题收集跟踪和反馈"],
    ["项目运营", "数城", "-", "日常运营工作", "系统需求收集跟踪和反馈"],
    ["项目运营", "数城", "-", "日常运营工作", "第三方平台或系统沟通协调"],
    ["项目运营", "数城", "-", "日常运营工作", "系统培训服务"],
    ["项目运营", "数城", "-", "验收工作", "验收工作"],
    ["项目运维", "中天", "-", "日常运维工作", "系统日常监控巡检"],
    ["项目运维", "中天", "-", "日常运维工作", "应用接入技术对接"],
    ["项目运维", "中天", "-", "日常运维工作", "系统故障分析和处理"],
    ["项目运维", "中天", "-", "日常运维工作", "系统版本升级发布"],
    ["项目运维", "中天", "-", "日常运维工作", "安全漏洞修复"],
    ["项目运维", "中天", "-", "验收工作", "验收工作"],
]
replace_first_table(doc, table_rows)

doc.save(OUT)
print(OUT)
